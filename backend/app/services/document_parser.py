import os
import re
import pypdf
import docx
import pandas as pd
from typing import List, Dict, Any

def is_garbled_line(line: str) -> bool:
    """
    Detects if a line is likely garbled OCR / font glyph extraction noise
    (e.g., 'liorx 1 DY C Fe 9 j xz 3 8 O C7 r Z r Z...').
    """
    tokens = line.strip().split()
    if not tokens or len(tokens) < 4:
        return False
    # If majority of words in a line are single characters or numbers, it's garbled font noise
    single_char_count = sum(1 for t in tokens if len(t) <= 1)
    if single_char_count / len(tokens) > 0.45 and len(tokens) > 6:
        return True
    return False

def normalize_text(text: str) -> str:
    """
    Strips boilerplate whitespace, normalizes special characters,
    filters garbled PDF font noise, and maintains paragraph boundaries.
    """
    # Replace multiple newlines with single newline
    text = re.sub(r'[\r\n]+', '\n', text)
    # Replace multiple spaces/tabs with single space
    text = re.sub(r'[ \t]+', ' ', text)
    
    cleaned_lines = []
    for line in text.split('\n'):
        if not is_garbled_line(line):
            cleaned_lines.append(line)
            
    return "\n".join(cleaned_lines).strip()

def parse_pdf(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text page-by-page from a PDF file.
    """
    pages = []
    with open(file_path, 'rb') as f:
        reader = pypdf.PdfReader(f)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            normalized = normalize_text(text)
            if normalized:
                pages.append({
                    "page_number": i + 1,
                    "text": normalized
                })
    return pages

def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text from DOCX files, including headings, paragraphs, and tables.
    """
    doc = docx.Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text.strip())
    
    # Extract tables and format rows
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))

    normalized = normalize_text("\n".join(full_text))
    # Word documents don't have standard physical pages, treat as page 1
    return [{"page_number": 1, "text": normalized}]

def parse_txt(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text from TXT files with UTF-8 normalization.
    """
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()
    normalized = normalize_text(text)
    return [{"page_number": 1, "text": normalized}]

def parse_csv(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts tabular data from CSV using pandas and formats each row into a text summary.
    """
    # Detect delimiter (comma or tab)
    try:
        df = pd.read_csv(file_path)
    except Exception:
        df = pd.read_csv(file_path, sep='\t')
        
    row_summaries = []
    for idx, row in df.iterrows():
        items = []
        for col in df.columns:
            val = row[col]
            if pd.notna(val):
                items.append(f"{col}: {val}")
        if items:
            row_summaries.append(f"Row {idx + 1}: " + ", ".join(items))
            
    normalized = normalize_text("\n".join(row_summaries))
    return [{"page_number": 1, "text": normalized}]

def extract_text_from_file(file_path: str) -> List[Dict[str, Any]]:
    """
    Dispatches file to correct extractor based on extension.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at: {file_path}")
        
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return parse_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return parse_docx(file_path)
    elif ext == '.txt':
        return parse_txt(file_path)
    elif ext == '.csv':
        return parse_csv(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
